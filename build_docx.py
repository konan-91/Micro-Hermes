#!/usr/bin/env python3
"""
Mechanical Markdown -> Word converter for dissertation.md.

Does NOT rewrite, summarize, or regenerate any prose. It only:
  - splits the source into paragraph blocks on blank lines,
  - classifies each block as a heading (exact string match against the
    document's own Table of Contents listing), a code block (4+ space
    indented run), or a plain paragraph,
  - writes each block's ORIGINAL text into a python-docx paragraph/run
    unchanged, applying only visual styling (bold/colour/font).

Run verify_docx.py afterwards to confirm every character of body text in
the .md made it into the .docx untouched.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(__file__).parent / "dissertation.md"
OUT = Path(__file__).parent / "dissertation.docx"

LIGHT_BLUE = RGBColor(0x4A, 0x90, 0xD9)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

CHAPTER_RE = re.compile(r"^\d+\.\s")
SECTION_RE = re.compile(r"^\d+\.\d+\s")
APPENDIX_RE = re.compile(r"^Appendix [A-Z]:")

# Headings that appear exactly once in the source (front-matter labels).
SINGLE_HEADINGS = {"Abstract", "Declaration", "Table of Contents"}


def classify_heading_level(text: str) -> int | None:
    """Return 1 for chapter-level, 2 for section-level, else None."""
    if text in SINGLE_HEADINGS or text == "References" or APPENDIX_RE.match(text):
        return 1
    if CHAPTER_RE.match(text):
        return 1
    if SECTION_RE.match(text):
        return 2
    return None


def read_blocks(lines: list[str]):
    """Yield (kind, text_or_lines) blocks, splitting on blank lines.

    kind is 'code' for a run of 4+-space-indented lines (merged, original
    text preserved line by line), otherwise 'line' for a single line.
    """
    i = 0
    n = len(lines)
    while i < n:
        if lines[i].strip() == "":
            i += 1
            continue
        run_start = i
        while i < n and lines[i].strip() != "":
            i += 1
        run = lines[run_start:i]
        if all(l.startswith("    ") for l in run) and len(run) > 0:
            yield ("code", run)
        else:
            for l in run:
                yield ("line", l)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and choose Update Field to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)


def set_update_fields_on_open(document):
    settings = document.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def main():
    raw_lines = SRC.read_text(encoding="utf-8").splitlines()

    # --- Front matter destined for the title page (pulled out of the body,
    # per the user's explicit request for a separate title page). These are
    # the first four non-blank lines of the source file.
    front_matter = []
    body_start = 0
    for idx, l in enumerate(raw_lines):
        if l.strip() != "":
            front_matter.append(l)
        if len(front_matter) == 4 and l.strip() != "":
            body_start = idx + 1
            break
    assert len(front_matter) == 4, front_matter
    title_text, subtitle_text, student_text, university_text = front_matter
    body_lines = raw_lines[body_start:]

    # --- Build ordered list of heading strings that appear TWICE (once in
    # the manual "Table of Contents" listing, once as the real heading).
    toc_idx = next(i for i, l in enumerate(body_lines) if l.strip() == "Table of Contents")
    # The manual listing runs from just after "Table of Contents" until the
    # line "1. Introduction" reappears a second time.
    double_heading_texts = []
    j = toc_idx + 1
    while True:
        line = body_lines[j].strip()
        if line != "":
            double_heading_texts.append(line)
            if line == "References":
                break
        j += 1

    seen_count = {t: 0 for t in double_heading_texts}

    document = Document()

    # ---------------- Title page ----------------
    section = document.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for _ in range(6):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DARK

    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle_text)
    r.font.size = Pt(15)
    r.font.color.rgb = LIGHT_BLUE

    for _ in range(4):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(student_text).font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(university_text).font.size = Pt(12)

    document.add_page_break()

    # ---------------- Table of contents page ----------------
    p = document.add_paragraph()
    r = p.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK

    document.add_paragraph()
    toc_anchor_p = document.add_paragraph()
    add_toc_field(toc_anchor_p)

    document.add_page_break()

    # ---------------- Body ----------------
    for kind, payload in read_blocks(body_lines):
        if kind == "code":
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for idx, code_line in enumerate(payload):
                if idx > 0:
                    p.add_run().add_break()
                r = p.add_run(code_line)
                r.font.name = "Courier New"
                r.font.size = Pt(10)
            continue

        line = payload
        text = line.strip()
        if text == "":
            continue

        level = None
        if text in SINGLE_HEADINGS:
            level = 1
        elif text in seen_count:
            seen_count[text] += 1
            if seen_count[text] == 2:
                level = classify_heading_level(text)
            else:
                level = None  # first occurrence: it's inside the manual ToC listing

        if level == 1:
            p = document.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = DARK
        elif level == 2:
            p = document.add_paragraph(style="Heading 2")
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = LIGHT_BLUE
        else:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line)
            r.font.size = Pt(11)

    set_update_fields_on_open(document)
    document.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    sys.exit(main())
